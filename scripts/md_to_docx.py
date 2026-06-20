# -*- coding: utf-8 -*-
"""Convierte archivos .md a .docx usando python-docx.
Uso: python scripts/md_to_docx.py [ruta_docs/]
Por defecto convierte todos los .md en docs/ al mismo directorio como .docx.
"""
import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_bg(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def add_run_with_inline_md(para, text):
    """Añade texto con soporte de **bold**, *italic* y `code`."""
    pattern = re.compile(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)')
    parts = pattern.split(text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = para.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = para.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        else:
            if part:
                para.add_run(part)


def parse_table(doc, lines, start):
    """Parsea una tabla markdown y la agrega al documento."""
    rows_data = []
    i = start
    while i < len(lines):
        line = lines[i].strip()
        if not line.startswith('|'):
            break
        # Ignorar separadores |---|---|
        cells = [c.strip() for c in line.strip('|').split('|')]
        if all(re.match(r'^[-: ]+$', c) for c in cells):
            i += 1
            is_header_sep = True
            continue
        rows_data.append(cells)
        i += 1

    if not rows_data:
        return i

    ncols = max(len(r) for r in rows_data)
    table = doc.add_table(rows=len(rows_data), cols=ncols)
    table.style = 'Table Grid'

    for ri, row in enumerate(rows_data):
        for ci, cell_text in enumerate(row):
            if ci >= ncols:
                break
            cell = table.cell(ri, ci)
            cell.text = ''
            para = cell.paragraphs[0]
            add_run_with_inline_md(para, cell_text)
            if ri == 0:
                for run in para.runs:
                    run.bold = True
                set_cell_bg(cell, 'D9E1F2')
    return i


def convert_md_to_docx(md_path: Path, out_path: Path):
    doc = Document()

    # Estilos de párrafo básicos
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)

    with open(md_path, encoding='utf-8') as f:
        lines = f.readlines()

    in_code_block = False
    code_lines = []
    i = 0

    while i < len(lines):
        line = lines[i].rstrip('\n')

        # Bloque de código
        if line.strip().startswith('```'):
            if in_code_block:
                # Cerrar bloque
                para = doc.add_paragraph()
                for cl in code_lines:
                    run = para.add_run(cl + '\n')
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x24, 0x29, 0x2E)
                para.paragraph_format.left_indent = Inches(0.3)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        stripped = line.strip()

        # Línea horizontal
        if re.match(r'^-{3,}$|^\*{3,}$|^_{3,}$', stripped):
            doc.add_paragraph('─' * 60)
            i += 1
            continue

        # Tabla
        if stripped.startswith('|'):
            i = parse_table(doc, lines, i)
            continue

        # Encabezados
        m = re.match(r'^(#{1,6})\s+(.*)', stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2)
            heading = doc.add_heading('', level=min(level, 4))
            add_run_with_inline_md(heading, text)
            i += 1
            continue

        # Lista con guión
        m = re.match(r'^-\s+(.*)', stripped)
        if m:
            para = doc.add_paragraph(style='List Bullet')
            add_run_with_inline_md(para, m.group(1))
            i += 1
            continue

        # Lista numerada
        m = re.match(r'^\d+\.\s+(.*)', stripped)
        if m:
            para = doc.add_paragraph(style='List Number')
            add_run_with_inline_md(para, m.group(1))
            i += 1
            continue

        # Línea de bloque de código indentado (4 espacios o tab)
        if line.startswith('    ') or line.startswith('\t'):
            para = doc.add_paragraph()
            run = para.add_run(line.lstrip())
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x24, 0x29, 0x2E)
            para.paragraph_format.left_indent = Inches(0.3)
            i += 1
            continue

        # Línea vacía
        if not stripped:
            i += 1
            continue

        # Párrafo normal
        para = doc.add_paragraph()
        add_run_with_inline_md(para, stripped)
        i += 1

    doc.save(out_path)
    print(f"  OK  {md_path.name}  ->  {out_path.name}")


def main():
    base = Path(__file__).resolve().parent.parent
    docs_dir = Path(sys.argv[1]) if len(sys.argv) > 1 else base / "docs"

    md_files = list(docs_dir.glob("*.md"))
    if not md_files:
        print(f"No se encontraron archivos .md en {docs_dir}")
        return

    print(f"Convirtiendo {len(md_files)} archivos en {docs_dir}/ ...")
    for md in sorted(md_files):
        out = md.with_suffix('.docx')
        convert_md_to_docx(md, out)

    print(f"\nListo. Archivos .docx guardados en {docs_dir}/"  )  # noqa


if __name__ == "__main__":
    main()
